
from docx import Document

from dojo.models import Finding


class PhilipsSecurityTestReportParser:

    """Security testing report for Philips devices from Security Center of Excellence (SCoE)"""

    @staticmethod
    def extract_vulnerability_report(doc_path, start_title):
        # Load the Word document
        doc = Document(doc_path)

        # Initialize variables
        extracting = False
        report_data = {
            "paragraphs": [],
            "tables": [],
        }

        # Iterate through paragraphs in the document to find the start title
        for para in doc.paragraphs:
            if start_title in para.text:
                extracting = True
                break

        # If the start title is found, iterate through tables in the document
        if extracting:
            for table in doc.tables:
                # Check if the first cell of the table contains "Vulnerability Title / Vulnerability Name"
                cell_text = table.cell(0, 0).text
                if cell_text == "Vulnerability Title" or cell_text == "Vulnerability Name":
                    table.cell(0, 0).text = "Vulnerability Title"    # Normalize the cell text, SCoE reports have different names
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    report_data["tables"].append(table_data)

        return report_data

    def get_scan_types(self):
        return ["Philips Security Test Report Scan"]

    def get_label_for_scan_types(self, scan_type):
        return "Philips Security Test Report Scan"

    def get_description_for_scan_types(self, scan_type):
        return "Philips Security Test Report report file can be imported in docx format."

    def get_findings(self, file, test):
        start_title = "Detailed Vulnerability Report"
        data = self.extract_vulnerability_report(file, start_title)
        findings = []

        for vulnerability in data["tables"]:
            vulnerability_info = {}

            # Process each key-value pair in the vulnerability
            for item in vulnerability:
                key = item[0]
                value = item[1]
                if key == "CVSS base Score/ Vector":
                    lines = value.split("\n")
                    cvss_base_score = lines[0].replace("CVSS Base Score: ", "").strip()
                    cvss_vector = lines[1].replace("CVSS Vector: ", "").strip()
                    vulnerability_info["CVSS Base Score"] = cvss_base_score
                    vulnerability_info["CVSS Vector"] = cvss_vector
                else:
                    vulnerability_info[key] = value

            find = Finding(
                title=vulnerability_info["Vulnerability Title"],
                test=test,
                description=vulnerability_info["Description"],
                severity=vulnerability_info["Severity"],
                mitigation=vulnerability_info["Recommendation"],
          #      cvssv3=CVSS3(vulnerability_info["CVSS Vector"]).clean_vector(),
          #      cvssv3_score=vulnerability_info["CVSS Base Score"],
                static_finding=True,
            )
            findings.append(find)

        return findings

    def convert_severity(self, num_severity):
        """Convert severity value"""
        if num_severity >= -10:
            return "Low"
        if -11 >= num_severity > -26:
            return "Medium"
        if num_severity <= -26:
            return "High"
        return "Info"
